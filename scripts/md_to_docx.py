#!/usr/bin/env python3
"""Convert ingestion-audit-kept-vs-dropped.md → .docx (fallback when TS docx is locked)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "phase-1-baseline-audit" / "ingestion-audit-kept-vs-dropped.md"
DOCX_PATH = ROOT / "docs" / "phase-1-baseline-audit" / "ingestion-audit-kept-vs-dropped.docx"


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    headers = [c.strip() for c in lines[start].strip("|").split("|")]
    rows: list[list[str]] = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip("|").split("|")])
        i += 1
    return headers, rows, i


def md_to_docx(md: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_buf))
                p.style = "Intense Quote"
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-:\s|]+\|$", lines[i + 1].strip()):
            headers, rows, i = parse_table(lines, i)
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                table.rows[0].cells[ci].text = h
                for p in table.rows[0].cells[ci].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            for ri, row in enumerate(rows):
                padded = row + [""] * (len(headers) - len(row))
                for ci, val in enumerate(padded[: len(headers)]):
                    table.rows[ri + 1].cells[ci].text = val
                    for p in table.rows[ri + 1].cells[ci].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)
            doc.add_paragraph()
            continue
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip().startswith("> "):
            doc.add_paragraph(line.strip()[2:], style="Intense Quote")
        elif line.strip() == "---":
            doc.add_paragraph("")
        elif line.strip():
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text)

        i += 1

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    return doc


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"Missing {MD_PATH} — run audit:ingestion-report first")
    md = MD_PATH.read_text(encoding="utf-8")
    doc = md_to_docx(md)
    try:
        doc.save(DOCX_PATH)
        print(f"Wrote {DOCX_PATH}")
    except PermissionError:
        alt = DOCX_PATH.with_name("ingestion-audit-kept-vs-dropped.generated.docx")
        doc.save(alt)
        print(f"Original locked — wrote {alt}")


if __name__ == "__main__":
    main()
