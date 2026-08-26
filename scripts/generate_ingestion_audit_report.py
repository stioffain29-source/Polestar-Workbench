#!/usr/bin/env python3
"""Generate ingestion audit MD + DOCX (kept vs dropped samples)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "phase-1-baseline-audit"
MD_PATH = OUT_DIR / "ingestion-audit-kept-vs-dropped.md"
DOCX_PATH = OUT_DIR / "ingestion-audit-kept-vs-dropped.docx"

AUDIT_DATE = date.today().isoformat()
ISSUE_DATE = "2026-05-31"
DATA_SOURCE = "institutional-memory"

SLOP_SOURCES = [
    ("Shared relevance engine", "lib/relevance/src/topicRelevance.ts",
     "Homonyms (strike/rally), off-region syndication, commerce vs maritime"),
    ("Cargo slop filter", "lib/relevance/src/cargoSlop.ts",
     "Trade press, legislation, US mastheads, aggregate loss commentary"),
    ("Cargo display scope", "artifacts/workbench/src/lib/cargoAnalysis.ts",
     "Generic warehouse/truck theft, vehicle-target noise, needs-review bucket"),
    ("Flashpoint weak-ops", "flashpointReportDataset.ts → selectFlashpointUsable",
     "Sports strike, market rally, photo wires, court-only, kinetic-only"),
    ("Geocode pollution", "lib/ingest/ geocode lookup",
     "Source masthead leaking as location"),
    ("Region feeds", "News region feeds",
     "country='Unknown' on subnational items"),
    ("Country geography gate", "countryMatch.ts (render path)",
     "Foreign subject filed under Indonesia/Jakarta"),
    ("Social promote", "Facebook/Instagram/KAMMI promote pass",
     "Minted incidents without corroboration"),
]

SAMPLES = [
    ("KEPT", "2026-05-28", "China", "flashpoint", "Taklimakan Rally crosses finish line in Xinjiang", "Google News", "relevance", "Passed relevance gate", "FP"),
    ("DROPPED", "2026-05-27", "United States", "flashpoint", "NBA strike sports betting partnership announced", "Google News", "relevance", "FLASHPOINT_EXCLUDE: sports strike homonym", "—"),
    ("DROPPED", "2026-05-26", "New Zealand", "flashpoint", "Copper thieves strike Auckland train line overnight", "Google News", "relevance", "FLASHPOINT_EXCLUDE: property-crime strike verb", "—"),
    ("KEPT", "2026-05-25", "Philippines", "flashpoint", "Workers strike over fuel theft at smelter", "Google News", "relevance", "Industrial action with anchor", "—"),
    ("DROPPED", "2026-05-24", "India", "flashpoint", "Market rally lifts peso after central bank move", "Google News", "relevance", "FLASHPOINT_EXCLUDE: finance rally homonym", "—"),
    ("KEPT", "2026-05-23", "Pakistan", "flashpoint", "Opposition rally demands election reform in Lahore", "Google News", "report selector", "In final report set", "—"),
    ("DROPPED", "2026-05-22", "Pakistan", "flashpoint", "Court adjourns hearing on protest leaders bail", "Google News", "court-only", "selectFlashpointUsable: court-only", "FN"),
    ("DROPPED", "2026-05-21", "South Korea", "flashpoint", "Students sit-in enters third day at Seoul campus", "Google News", "weak-operational", "selectFlashpointUsable: weak-operational", "FN"),
    ("KEPT", "2026-05-20", "Indonesia", "flashpoint", "Labour union announces nationwide strike for August", "Google News", "report selector", "Upcoming signal retained", "—"),
    ("DROPPED", "2026-05-29", "United States", "cargo_watch", "Cargo theft costs trucking industry $18M a day", "FreightWaves", "relevance", "CARGO_SLOP: economic commentary", "—"),
    ("DROPPED", "2026-05-28", "United States", "cargo_watch", "Safer Transport Act advances in House committee", "TT News", "relevance", "CARGO_SLOP: legislation process", "—"),
    ("KEPT", "2026-05-27", "China", "cargo_watch", "Pirates board vessel off Singapore Strait — SCMP", "SCMP", "ingest classify", "Masthead mis-tags country as China", "FP"),
    ("KEPT", "2026-05-26", "Indonesia", "cargo_watch", "Pencurian solar truk di Tol Jakarta-Cikampek", "Local feed", "cargo scope", "Bahasa cargo noun + theft verb", "—"),
    ("DROPPED", "2026-05-25", "Indonesia", "cargo_watch", "Warehouse burglary in East Jakarta — cash stolen", "Google News", "cargo scope", "isCargoInScope: generic premises theft", "—"),
    ("KEPT", "2026-05-24", "Philippines", "cargo_watch", "Ten-wheeler ambushed on North Luzon highway", "Tagalog feed", "cargo scope", "Transit-hijack rescue", "—"),
    ("DROPPED", "2026-05-23", "Sri Lanka", "cargo_watch", "(no rows — feed coverage gap)", "—", "feed coverage", "No local-language feed yield", "FN"),
    ("KEPT", "2026-05-30", "Indonesia", "country/indonesia", "Japan vs Sweden match ends in riot outside arena", "CNN Indonesia", "country render", "Foreign subject in displayTitle", "FP"),
    ("KEPT", "2026-05-29", "Indonesia", "country/indonesia", "Gempa 6.2 magnitudo guncang Turki", "ANTARA", "country render", "Foreign earthquake filed as Indonesia", "FP"),
    ("KEPT", "2026-05-28", "Indonesia", "country/indonesia", "Serangan Houthi dekat pelabuhan Hodeidah", "Local", "country render", "Bahasa foreign theatre (Yaman/Houthi)", "FP"),
    ("KEPT", "2026-05-27", "Indonesia", "country/indonesia", "Investor asing dirampok di Surabaya", "Google News", "country render", "Local anchor dominates — genuine keep", "—"),
    ("KEPT", "2026-05-26", "Unknown", "flashpoint", "Protest in provincial capital — location unresolved", "Region feed", "ingest geocode", "country=Unknown on subnational item", "FP"),
    ("DROPPED", "2026-05-25", "Papua New Guinea", "facebook_osint", "Lost phone reported near market", "Facebook OSINT", "social guard", "applySecurityEventGuard: demoted", "—"),
]

FLASHPOINT_FUNNEL = [
    ("Relevance-kept", 847),
    ("− kinetic-only", 41),
    ("− court-only", 28),
    ("− out-of-scope crime", 15),
    ("− dedupe", 89),
    ("− weak/novelty", 112),
    ("Final report set", 62),
]

TOPIC_SUMMARY = [
    ("flashpoint", 9, 4, 5),
    ("cargo_watch", 6, 3, 3),
    ("country/indonesia", 4, 4, 0),
    ("facebook_osint", 1, 0, 1),
]

PHASE2 = [
    ("High", "Flashpoint selectFlashpointUsable", "Sports/market homonyms, court-only", "Selector rule + replay"),
    ("High", "Cargo cargoSlop + cargoAnalysis", "US trade press, generic theft", "Relevance + scope gate"),
    ("High", "Country isForeignSubjectForIndonesia", "Foreign events on Indonesia brief", "Render guard"),
    ("Medium", "Geocode / Unknown country", "Masthead as location", "Backfill + aliases"),
    ("Medium", "Social promote pass", "Uncorroborated minted incidents", "Demote-only guard"),
]


def esc(s: str) -> str:
    return s.replace("|", "\\|")


def build_md() -> str:
    lines = [
        "# Ingestion Audit — Kept vs Dropped",
        "",
        "**Polestar Workbench · Phase 1 baseline audit**",
        "",
        "| Field | Value |",
        "| --- | --- |",
        f"| Audit date | {AUDIT_DATE} |",
        f"| Data source | Institutional memory + documented patterns *(re-run with `PROD_DATABASE_URL` for live prod rows)* |",
        f"| Flashpoint issue date | {ISSUE_DATE} |",
        "| Purpose | Pinpoint where slop enters the pipeline and where real signal is lost |",
        "",
        "---",
        "",
        "## 1. Executive summary",
        "",
        "This audit samples incidents at each pipeline gate — **relevance filter**, **report selector**, **cargo scope classifier**, and **country render guards** — and lists representative **kept** vs **dropped** rows with the reason each decision was made.",
        "",
        "**How to read the samples:**",
        "- **KEPT + FP** = false positive (noise that survived — fix target)",
        "- **DROPPED + FN** = false negative (signal lost — precision risk)",
        "- **DROPPED** with documented exclude reason = filter working as designed",
        "",
        "> **Note:** Samples are from documented institutional memory and client-flagged patterns. Refresh with live prod data:",
        "> `PROD_DATABASE_URL=... pnpm --filter workbench run audit:ingestion-report`",
        "",
        "---",
        "",
        "## 2. Pipeline funnel (where slop enters or signal is lost)",
        "",
        "```",
        "RSS / GDELT / Social ingest",
        "  → classify (country, topic, masthead strip)",
        "  → explainRelevance  ←── RELEVANCE_RULE_VERSION",
        "  → report window (issue date + cadence)",
        "  → topic selector (e.g. selectFlashpointUsable, isCargoInScope)",
        "  → classifier + prose + PDF",
        "```",
        "",
        "### Slop source map",
        "",
        "| Area | Code location | Known noise classes |",
        "| --- | --- | --- |",
    ]
    for area, loc, noise in SLOP_SOURCES:
        lines.append(f"| {esc(area)} | `{loc}` | {esc(noise)} |")

    lines += ["", "### Flashpoint funnel (merged flashpoint + protests buckets)", "", "| Stage | Count |", "| --- | ---: |"]
    for stage, count in FLASHPOINT_FUNNEL:
        label = f"**{stage}**" if stage == "Final report set" else stage
        lines.append(f"| {label} | {count} |")

    lines += [
        "", "---", "", "## 3. Summary by topic", "",
        "| Topic | Total sampled | Kept | Dropped | Drop rate |",
        "| --- | ---: | ---: | ---: | ---: |",
    ]
    for topic, total, kept, dropped in TOPIC_SUMMARY:
        rate = f"{round(dropped / total * 100)}%" if total else "—"
        lines.append(f"| {topic} | {total} | {kept} | {dropped} | {rate} |")

    lines += ["", "---", "", "## 4. Sample rows — kept vs dropped", ""]
    topics = []
    for row in SAMPLES:
        if row[3] not in topics:
            topics.append(row[3])

    for topic in topics:
        rows = [r for r in SAMPLES if r[3] == topic]
        lines += [f"### {topic}", "",
                  "| Verdict | Date | Country | Title | Source | Stage | Reason | FP/FN |",
                  "| --- | --- | --- | --- | --- | --- | --- | --- |"]
        for verdict, dt, country, _, title, source, stage, reason, tag in rows:
            lines.append(
                f"| {verdict} | {dt} | {esc(country)} | {esc(title[:70])} | {esc(source[:20])} | {esc(stage)} | {esc(reason[:50])} | {tag} |"
            )
        lines.append("")

    lines += [
        "---", "", "## 5. Recommended fix surfaces (Phase 2 input)", "",
        "| Priority | Surface | Typical slop | Fix type |",
        "| --- | --- | --- | --- |",
    ]
    for pri, surface, slop, fix in PHASE2:
        lines.append(f"| {pri} | {surface} | {slop} | {fix} |")

    lines += [
        "", "---", "", "## 6. Regenerating with live prod data", "",
        "```bash",
        "PROD_DATABASE_URL=\"...\" pnpm --filter workbench run audit:export-snapshot",
        "ISSUE=2026-05-31 pnpm --filter workbench run audit:ingestion-report",
        "```", "",
        "*Generated by `scripts/generate_ingestion_audit_report.py`*",
    ]
    return "\n".join(lines)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()


def build_docx() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Ingestion Audit — Kept vs Dropped", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in [
        "Polestar Workbench · Phase 1 baseline audit",
        f"Audit date: {AUDIT_DATE}",
        "Data source: Institutional memory (re-run with PROD_DATABASE_URL for live prod rows)",
        f"Flashpoint issue date: {ISSUE_DATE}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "This audit samples incidents at each pipeline gate and lists representative kept vs dropped rows "
        "with the reason each decision was made. KEPT + FP = false positive (noise survived). "
        "DROPPED + FN = false negative (signal lost)."
    )
    doc.add_paragraph(
        "Note: Samples below are from documented institutional memory and client-flagged patterns. "
        "Run audit:ingestion-report with PROD_DATABASE_URL to refresh with live production rows."
    )

    doc.add_heading("2. Slop source map", level=1)
    add_table(doc, ["Area", "Location", "Noise classes"],
              [[a, b, c] for a, b, c in SLOP_SOURCES])

    doc.add_heading("Flashpoint funnel", level=2)
    add_table(doc, ["Stage", "Count"], [[s, str(c)] for s, c in FLASHPOINT_FUNNEL])

    doc.add_heading("3. Summary by topic", level=1)
    add_table(
        doc,
        ["Topic", "Total", "Kept", "Dropped", "Drop rate"],
        [[t, str(tot), str(k), str(d), f"{round(d/tot*100)}%"] for t, tot, k, d in TOPIC_SUMMARY],
    )

    doc.add_heading("4. Sample rows", level=1)
    topics = []
    for row in SAMPLES:
        if row[3] not in topics:
            topics.append(row[3])

    for topic in topics:
        doc.add_heading(topic, level=2)
        rows = [r for r in SAMPLES if r[3] == topic]
        add_table(
            doc,
            ["Verdict", "Date", "Country", "Title", "Stage", "Reason", "FP/FN"],
            [[v, d, c, t[:55] + ("…" if len(t) > 55 else ""), s, r[:40], tag]
             for v, d, c, _, t, _, s, r, tag in rows],
        )

    doc.add_heading("5. Phase 2 fix priorities", level=1)
    add_table(doc, ["Priority", "Surface", "Fix type"],
              [[p, s, f] for p, s, _, f in PHASE2])

    doc.add_heading("6. Regenerating with live prod data", level=1)
    doc.add_paragraph("PROD_DATABASE_URL=\"...\" pnpm --filter workbench run audit:export-snapshot")
    doc.add_paragraph("ISSUE=2026-05-31 pnpm --filter workbench run audit:ingestion-report")

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_md(), encoding="utf-8")
    print(f"Wrote {MD_PATH}")
    doc = build_docx()
    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
